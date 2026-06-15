from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

from app.main import app
from app.generation.providers import TemplateLLMProvider, create_llm_provider
from app.ingestion.parsers import parse_supported_file
from app.retrieval.fulltext import NoopFullTextSearch
from app.retrieval.embeddings import HashingEmbedder, create_embedder


client = TestClient(app)


def auth_headers(username: str = "admin", password: str = "rag-console") -> dict[str, str]:
    response = client.post(
        "/api/v1/auth/login",
        json={"username": username, "password": password},
    )
    assert response.status_code == 200
    token = response.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_healthcheck() -> None:
    response = client.get("/health")
    assert response.status_code == 200
    assert response.json() == {"status": "ok"}

    db_response = client.get("/health/db")
    assert db_response.status_code == 200
    assert db_response.json()["status"] == "ok"

    config_response = client.get("/config", headers=auth_headers())
    assert config_response.status_code == 200
    assert config_response.json()["runtime_mode"] == "local"
    assert "preferences" in config_response.json()

    analytics_response = client.get("/analytics/overview")
    assert analytics_response.status_code == 200
    assert "document_count" in analytics_response.json()


def test_ingest_and_query() -> None:
    ingest_payload = {
        "document_id": "doc-001",
        "title": "RAG Design Spec",
        "source": "internal-wiki",
        "content": "RAG systems should support chunking, hybrid retrieval, reranking, and grounded answers.",
        "acl": ["engineering"],
        "metadata": {"department": "engineering"},
    }
    ingest_response = client.post("/api/v1/documents/ingest", json=ingest_payload)
    ingest_response = client.post("/api/v1/documents/ingest", json=ingest_payload, headers=auth_headers())
    assert ingest_response.status_code == 201

    query_payload = {
        "question": "What capabilities should a RAG system support?",
        "filters": {"department": "engineering"},
    }
    query_response = client.post("/api/v1/query", json=query_payload, headers=auth_headers("analyst", "rag-analyst"))
    assert query_response.status_code == 200
    body = query_response.json()
    assert body["citations"]
    assert body["confidence"] > 0
    assert "content_preview" in body["citations"][0]

    semantic_query = client.post(
        "/api/v1/query",
        json={
            "question": "What capabilities should a RAG system support?",
            "filters": {"department": "engineering"},
            "retrieval_mode": "semantic",
        },
    )
    assert semantic_query.status_code == 200

    logs_response = client.get("/api/v1/query/logs", headers=auth_headers("auditor", "rag-audit"))
    assert logs_response.status_code == 200
    logs = logs_response.json()
    assert logs
    assert any(item["user_id"] == "analyst" and item["role"] == "editor" for item in logs)


def test_hybrid_retrieval_uses_vector_signal() -> None:
    ingest_payload = {
        "document_id": "doc-002",
        "title": "Search Notes",
        "source": "engineering-notes",
        "content": "The platform uses reranking to improve retrieval quality after recall.",
        "acl": ["engineering"],
        "metadata": {"department": "engineering"},
    }
    response = client.post("/api/v1/documents/ingest", json=ingest_payload, headers=auth_headers())
    assert response.status_code == 201

    query_response = client.post(
        "/api/v1/query",
        json={
            "question": "How does rerank improve recall quality?",
            "filters": {"department": "engineering"},
        },
        headers=auth_headers("analyst", "rag-analyst"),
    )
    assert query_response.status_code == 200
    body = query_response.json()
    assert body["citations"]
    assert body["citations"][0]["document_id"] == "doc-002"


def test_chunk_acl_and_retrieval_debug() -> None:
    ingest_payload = {
        "document_id": "doc-acl",
        "title": "ACL Notes",
        "source": "security",
        "content": "finance secret budget policy should only be visible to finance users.",
        "acl": ["engineering"],
        "metadata": {"department": "security", "chunk_acl": "finance"},
    }
    response = client.post("/api/v1/documents/ingest", json=ingest_payload, headers=auth_headers())
    assert response.status_code == 201

    blocked = client.post(
        "/api/v1/query",
        json={
            "question": "finance secret budget policy",
            "filters": {"department": "security"},
        },
        headers=auth_headers("analyst", "rag-analyst"),
    )
    assert blocked.status_code == 200
    assert not blocked.json()["citations"]

    allowed = client.post(
        "/api/v1/query/debug",
        json={
            "question": "finance secret budget policy",
            "filters": {"department": "security"},
        },
        headers=auth_headers("auditor", "rag-audit"),
    )
    assert allowed.status_code == 200
    body = allowed.json()
    assert body["chunks"]
    assert body["chunks"][0]["document_id"] == "doc-acl"
    assert body["chunks"][0]["lexical_score"] >= 0
    assert body["chunks"][0]["semantic_score"] >= 0


def test_sync_local_directory(tmp_path: Path) -> None:
    docs_dir = tmp_path / "docs"
    docs_dir.mkdir()
    (docs_dir / "architecture.md").write_text(
        "# Architecture\nHybrid retrieval combines keyword search and dense retrieval.",
        encoding="utf-8",
    )
    (docs_dir / "faq.txt").write_text(
        "Audit logs should capture queries, answers, and citations.",
        encoding="utf-8",
    )

    response = client.post(
        "/api/v1/documents/sync/local-dir",
        json={
            "root_path": str(docs_dir),
            "recursive": True,
            "extensions": [".md", ".txt"],
            "default_acl": ["engineering"],
            "default_metadata": {"department": "engineering"},
            "source_label": "local-sync",
            "max_files": 20,
        },
        headers=auth_headers(),
    )

    assert response.status_code == 200
    body = response.json()
    assert body["imported_documents"] == 2
    assert body["scanned_files"] == 2
    assert not body["skipped_reasons"]


def test_document_management_endpoints() -> None:
    payload = {
        "document_id": "doc-manage",
        "title": "Managed Document",
        "source": "manual",
        "content": "This document is used to test management endpoints.",
        "acl": ["engineering"],
        "metadata": {"department": "engineering"},
    }
    ingest_response = client.post("/api/v1/documents/ingest", json=payload, headers=auth_headers())
    assert ingest_response.status_code == 201

    detail_response = client.get("/api/v1/documents/doc-manage", headers=auth_headers("analyst", "rag-analyst"))
    assert detail_response.status_code == 200
    assert detail_response.json()["document_id"] == "doc-manage"

    chunks_response = client.get("/api/v1/documents/doc-manage/chunks", headers=auth_headers("analyst", "rag-analyst"))
    assert chunks_response.status_code == 200
    assert chunks_response.json()

    delete_response = client.delete("/api/v1/documents/doc-manage", headers=auth_headers())
    assert delete_response.status_code == 204

    missing_response = client.get("/api/v1/documents/doc-manage", headers=auth_headers("analyst", "rag-analyst"))
    assert missing_response.status_code == 404


def test_document_update_and_batch_delete_endpoints() -> None:
    first = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "doc-update-1",
            "title": "Original Title One",
            "source": "manual",
            "content": "first document for update and batch delete",
            "acl": [],
            "metadata": {"category": "draft"},
        },
        headers=auth_headers(),
    )
    second = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "doc-update-2",
            "title": "Original Title Two",
            "source": "manual",
            "content": "second document for update and batch delete",
            "acl": [],
            "metadata": {"category": "draft"},
        },
        headers=auth_headers(),
    )
    assert first.status_code == 201
    assert second.status_code == 201

    updated = client.patch(
        "/api/v1/documents/doc-update-1",
        json={
            "title": "Renamed Knowledge Doc",
            "source": "knowledge-base",
            "metadata": {"category": "product", "owner": "ops"},
        },
        headers=auth_headers(),
    )
    assert updated.status_code == 200
    body = updated.json()
    assert body["title"] == "Renamed Knowledge Doc"
    assert body["source"] == "knowledge-base"
    assert body["metadata"]["category"] == "product"
    assert body["metadata"]["owner"] == "ops"

    detail = client.get("/api/v1/documents/doc-update-1", headers=auth_headers("analyst", "rag-analyst"))
    assert detail.status_code == 200
    detail_body = detail.json()
    assert detail_body["title"] == "Renamed Knowledge Doc"
    assert detail_body["metadata"]["category"] == "product"

    batch_delete = client.post(
        "/api/v1/documents/batch-delete",
        json={"document_ids": ["doc-update-1", "doc-update-2"]},
        headers=auth_headers(),
    )
    assert batch_delete.status_code == 200
    batch_body = batch_delete.json()
    assert batch_body["deleted_count"] == 2
    assert batch_body["missing_document_ids"] == []

    remaining_first = client.get("/api/v1/documents/doc-update-1", headers=auth_headers("analyst", "rag-analyst"))
    remaining_second = client.get("/api/v1/documents/doc-update-2", headers=auth_headers("analyst", "rag-analyst"))
    assert remaining_first.status_code == 404
    assert remaining_second.status_code == 404


def test_document_batch_update_and_move_endpoints(tmp_path: Path) -> None:
    first = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "doc-batch-1",
            "title": "Batch One",
            "source": "manual",
            "content": "first document for batch category update",
            "acl": [],
            "metadata": {"category": "draft"},
        },
        headers=auth_headers(),
    )
    second = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "doc-batch-2",
            "title": "Batch Two",
            "source": "manual",
            "content": "second document for batch category update",
            "acl": [],
            "metadata": {"category": "draft"},
        },
        headers=auth_headers(),
    )
    assert first.status_code == 201
    assert second.status_code == 201

    batch_update = client.post(
        "/api/v1/documents/batch-update",
        json={
            "document_ids": ["doc-batch-1", "doc-batch-2"],
            "metadata_updates": {"category": "operations", "folder_path": "archive/2026"},
        },
        headers=auth_headers(),
    )
    assert batch_update.status_code == 200
    batch_body = batch_update.json()
    assert batch_body["updated_count"] == 2
    assert batch_body["missing_document_ids"] == []

    first_detail = client.get("/api/v1/documents/doc-batch-1", headers=auth_headers("analyst", "rag-analyst"))
    second_detail = client.get("/api/v1/documents/doc-batch-2", headers=auth_headers("analyst", "rag-analyst"))
    assert first_detail.status_code == 200
    assert second_detail.status_code == 200
    assert first_detail.json()["metadata"]["category"] == "operations"
    assert first_detail.json()["metadata"]["folder_path"] == "archive/2026"
    assert second_detail.json()["metadata"]["category"] == "operations"

    upload_path = tmp_path / "move-target.md"
    upload_path.write_text("# Move Target\nUploaded file for move endpoint.", encoding="utf-8")

    with upload_path.open("rb") as handle:
      uploaded = client.post(
            "/api/v1/documents/upload",
            headers=auth_headers(),
            files={"file": ("move-target.md", handle, "text/markdown")},
            data={
                "source": "ui-upload",
                "acl": "",
                "metadata_json": '{"category":"draft"}',
            },
        )

    assert uploaded.status_code == 201
    uploaded_id = uploaded.json()["document_id"]
    old_path = Path(uploaded.json()["file_path"])
    assert old_path.exists()

    moved = client.post(
        "/api/v1/documents/move",
        json={
            "document_ids": [uploaded_id],
            "folder_path": "archive/uploads",
        },
        headers=auth_headers(),
    )
    assert moved.status_code == 200
    moved_body = moved.json()
    assert moved_body["moved_count"] == 1
    assert moved_body["missing_document_ids"] == []

    moved_detail = client.get(f"/api/v1/documents/{uploaded_id}", headers=auth_headers("analyst", "rag-analyst"))
    assert moved_detail.status_code == 200
    moved_path = Path(moved_detail.json()["file_path"])
    assert moved_detail.json()["metadata"]["folder_path"] == "archive/uploads"
    assert moved_path.exists()
    assert moved_path != old_path
    assert not old_path.exists()


def test_document_rename_file_endpoint(tmp_path: Path) -> None:
    upload_path = tmp_path / "rename-target.md"
    upload_path.write_text("# Rename Target\nUploaded file for rename endpoint.", encoding="utf-8")

    with upload_path.open("rb") as handle:
        uploaded = client.post(
            "/api/v1/documents/upload",
            headers=auth_headers(),
            files={"file": ("rename-target.md", handle, "text/markdown")},
            data={
                "source": "ui-upload",
                "acl": "",
                "metadata_json": '{"category":"draft","folder_path":"archive/uploads"}',
            },
        )

    assert uploaded.status_code == 201
    document_id = uploaded.json()["document_id"]
    old_path = Path(uploaded.json()["file_path"])
    assert old_path.exists()

    renamed = client.post(
        "/api/v1/documents/rename-file",
        headers=auth_headers(),
        json={
            "document_id": document_id,
            "filename": "renamed-knowledge-note.md",
        },
    )
    assert renamed.status_code == 200
    renamed_body = renamed.json()
    new_path = Path(renamed_body["file_path"])
    assert renamed_body["metadata"]["file_name"] == "renamed-knowledge-note.md"
    assert new_path.name == "renamed-knowledge-note.md"
    assert new_path.exists()
    assert not old_path.exists()

    download = client.get(f"/api/v1/documents/{document_id}/file", headers=auth_headers("analyst", "rag-analyst"))
    assert download.status_code == 200
    assert "Uploaded file for rename endpoint." in download.text


def test_upload_document_endpoint(tmp_path: Path) -> None:
    upload_path = tmp_path / "uploaded.md"
    upload_path.write_text("# Upload\nUploaded content for knowledge base.", encoding="utf-8")

    with upload_path.open("rb") as handle:
        response = client.post(
            "/api/v1/documents/upload",
            headers=auth_headers(),
            files={"file": ("uploaded.md", handle, "text/markdown")},
            data={
                "source": "ui-upload",
                "acl": "engineering",
                "metadata_json": '{"department":"engineering"}',
            },
        )

    assert response.status_code == 201
    assert response.json()["title"] == "uploaded"
    assert response.json()["file_path"]
    document_id = response.json()["document_id"]

    download = client.get(f"/api/v1/documents/{document_id}/file", headers=auth_headers("analyst", "rag-analyst"))
    assert download.status_code == 200
    assert "Uploaded content" in download.text


def test_logs_require_auditor_or_admin() -> None:
    response = client.get("/api/v1/query/logs", headers=auth_headers("analyst", "rag-analyst"))
    assert response.status_code == 403


def test_logs_accessible_in_local_mode_without_auth() -> None:
    response = client.get("/api/v1/query/logs")
    assert response.status_code == 200


def test_refresh_token() -> None:
    headers = auth_headers("admin", "rag-console")
    response = client.post("/api/v1/auth/refresh", headers=headers)
    assert response.status_code == 200
    body = response.json()
    assert body["access_token"]
    assert body["user"]["username"] == "admin"


def test_log_filters_and_export() -> None:
    seed = client.post(
        "/api/v1/query",
        json={"question": "engineering knowledge", "filters": {"department": "engineering"}},
        headers=auth_headers("auditor", "rag-audit"),
    )
    assert seed.status_code == 200

    filtered = client.get(
        "/api/v1/query/logs",
        params={"user_id": "auditor", "role": "auditor", "question": "engineering"},
        headers=auth_headers("auditor", "rag-audit"),
    )
    assert filtered.status_code == 200
    body = filtered.json()
    assert body
    assert all(item["user_id"] == "auditor" for item in body)

    exported = client.get(
        "/api/v1/query/logs/export",
        params={"user_id": "auditor"},
        headers=auth_headers("auditor", "rag-audit"),
    )
    assert exported.status_code == 200
    assert "query-logs.csv" in exported.headers["content-disposition"]
    assert "user_id" in exported.text


def test_maintenance_rebuild_and_backup_restore() -> None:
    ingest_response = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "maint-doc",
            "title": "Maintenance Doc",
            "source": "ops",
            "content": "maintenance backup restore rebuild index workflow",
            "acl": [],
            "metadata": {"category": "operations"},
        },
        headers=auth_headers(),
    )
    assert ingest_response.status_code == 201

    rebuild = client.post("/maintenance/rebuild-indexes", headers=auth_headers())
    assert rebuild.status_code == 200
    rebuild_body = rebuild.json()
    assert rebuild_body["success"] is True
    assert rebuild_body["document_count"] >= 1

    exported = client.get("/maintenance/backup/export", headers=auth_headers())
    assert exported.status_code == 200
    assert "rag-backup.json" in exported.headers["content-disposition"]
    assert "maint-doc" in exported.text

    deleted = client.delete("/api/v1/documents/maint-doc", headers=auth_headers())
    assert deleted.status_code == 204

    restored = client.post(
        "/maintenance/backup/restore",
        headers=auth_headers(),
        files={"file": ("rag-backup.json", exported.text.encode("utf-8"), "application/json")},
    )
    assert restored.status_code == 200
    restore_body = restored.json()
    assert restore_body["success"] is True
    assert restore_body["document_count"] >= 1

    detail = client.get("/api/v1/documents/maint-doc", headers=auth_headers())
    assert detail.status_code == 200


def test_backup_version_list_and_verify() -> None:
    response = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "backup-version-doc",
            "title": "Backup Version Doc",
            "source": "backup-test",
            "content": "backup version verification content",
            "acl": [],
            "metadata": {"category": "maintenance"},
        },
        headers=auth_headers(),
    )
    assert response.status_code == 201

    created = client.post("/maintenance/backup/create", headers=auth_headers())
    assert created.status_code == 200
    created_body = created.json()
    assert created_body["filename"].endswith(".json")
    assert created_body["valid"] is True

    listed = client.get("/maintenance/backup/list", headers=auth_headers())
    assert listed.status_code == 200
    assert any(item["filename"] == created_body["filename"] for item in listed.json()["items"])

    verified = client.post(f"/maintenance/backup/verify/{created_body['filename']}", headers=auth_headers())
    assert verified.status_code == 200
    assert verified.json()["valid"] is True
    assert verified.json()["document_count"] >= 1


def test_maintenance_deduplicate_documents() -> None:
    first = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "dup-1",
            "title": "Duplicate One",
            "source": "ops",
            "content": "same content for duplicate cleanup",
            "acl": [],
            "metadata": {"category": "dup"},
        },
        headers=auth_headers(),
    )
    second = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "dup-2",
            "title": "Duplicate Two",
            "source": "ops",
            "content": "same content for duplicate cleanup",
            "acl": [],
            "metadata": {"category": "dup"},
        },
        headers=auth_headers(),
    )
    assert first.status_code == 201
    assert second.status_code == 201

    cleanup = client.post("/maintenance/cleanup-duplicates", headers=auth_headers())
    assert cleanup.status_code == 200
    body = cleanup.json()
    assert body["success"] is True
    assert body["document_count"] >= 1

    remaining = client.get("/api/v1/documents", headers=auth_headers())
    assert remaining.status_code == 200
    duplicate_docs = [item for item in remaining.json() if item["document_id"] in {"dup-1", "dup-2"}]
    assert len(duplicate_docs) == 1


def test_maintenance_cleanup_orphan_files(tmp_path: Path) -> None:
    orphan_dir = tmp_path / "uploads"
    orphan_dir.mkdir()
    orphan_file = orphan_dir / "orphan.txt"
    orphan_file.write_text("orphan payload", encoding="utf-8")

    response = client.post(
        "/maintenance/cleanup-orphans",
        headers=auth_headers(),
        json={"storage_root": str(orphan_dir)},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["success"] is True
    assert body["document_count"] >= 1
    assert not orphan_file.exists()


def test_retrieval_evaluation_endpoint_returns_hits() -> None:
    ingest = client.post(
        "/api/v1/documents/ingest",
        json={
            "document_id": "eval-doc",
            "title": "Evaluation Doc",
            "source": "eval",
            "content": "retrieval evaluation should return the relevant document first",
            "acl": [],
            "metadata": {"category": "evaluation"},
        },
        headers=auth_headers(),
    )
    assert ingest.status_code == 201

    response = client.post(
        "/api/v1/query/evaluate",
        headers=auth_headers(),
        json={
            "cases": [
                {
                    "question": "relevant document first",
                    "expected_document_id": "eval-doc",
                    "retrieval_mode": "hybrid",
                }
            ]
        },
    )
    assert response.status_code == 200
    body = response.json()
    assert body["cases"]
    assert body["cases"][0]["matched"] is True


def test_hashing_embedder_is_stable() -> None:
    embedder = HashingEmbedder(dimensions=64)
    left = embedder.embed_text("reranking improves recall quality")
    right = embedder.embed_text("reranking improves recall quality")
    assert left == right


def test_create_hashing_embedder() -> None:
    embedder = create_embedder(
        provider="hashing",
        dimensions=32,
        model="unused",
        api_key=None,
        base_url=None,
        timeout=30,
    )
    assert len(embedder.embed_text("hello")) == 32


def test_template_llm_provider() -> None:
    provider = create_llm_provider(
        provider="template",
        model="unused",
        api_key=None,
        base_url=None,
        timeout=30,
        temperature=0.2,
    )
    assert isinstance(provider, TemplateLLMProvider)


def test_noop_full_text_search() -> None:
    search = NoopFullTextSearch()
    assert search.search("rag", limit=5) == {}


def test_parse_docx_and_excel(tmp_path: Path) -> None:
    docx_path = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("Security Policy", level=1)
    document.add_paragraph("Employees must follow access control rules.")
    document.save(docx_path)

    title, content = parse_supported_file(docx_path)
    assert title == "policy"
    assert "access control" in content

    xlsx_path = tmp_path / "inventory.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Assets"
    sheet.append(["Name", "Owner"])
    sheet.append(["RAG Platform", "Engineering"])
    workbook.save(xlsx_path)

    title, content = parse_supported_file(xlsx_path)
    assert title == "inventory"
    assert "RAG Platform" in content
