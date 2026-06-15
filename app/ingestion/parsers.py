from __future__ import annotations

import json
import tempfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


TEXT_ENCODINGS = ("utf-8", "utf-8-sig", "gb18030")


def read_text_with_fallback(path: Path) -> str:
    for encoding in TEXT_ENCODINGS:
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def parse_supported_file(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        content = read_text_with_fallback(path)
        title = path.stem.replace("_", " ").replace("-", " ").strip() or path.name
        return title, content

    if suffix == ".json":
        raw = read_text_with_fallback(path)
        data = json.loads(raw)
        if isinstance(data, dict):
            title = str(data.get("title") or path.stem)
            if "content" in data:
                content = str(data["content"])
            else:
                content = json.dumps(data, ensure_ascii=False, indent=2)
            return title, content
        return path.stem, json.dumps(data, ensure_ascii=False, indent=2)

    if suffix == ".pdf":
        return _parse_pdf(path)

    if suffix == ".docx":
        return _parse_docx(path)

    if suffix in {".xlsx", ".xlsm"}:
        return _parse_excel(path)

    raise ValueError(f"Unsupported file type: {path.suffix}")


def parse_uploaded_bytes(filename: str, data: bytes) -> tuple[str, str]:
    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as handle:
        handle.write(data)
        temp_path = Path(handle.name)
    try:
        title, content = parse_supported_file(temp_path)
        fallback_title = temp_path.stem.replace("_", " ").replace("-", " ").strip()
        original_title = Path(filename).stem.replace("_", " ").replace("-", " ").strip() or Path(filename).name
        if title == fallback_title:
            title = original_title
        return title, content
    finally:
        temp_path.unlink(missing_ok=True)


def _parse_pdf(path: Path) -> tuple[str, str]:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[page {page_number}]\n{text.strip()}")
    content = "\n\n".join(pages).strip()
    if not content:
        raise ValueError("No extractable text found in PDF.")
    return path.stem, content


def _parse_docx(path: Path) -> tuple[str, str]:
    document = Document(str(path))
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    content = "\n".join(parts).strip()
    if not content:
        raise ValueError("No extractable text found in Word document.")
    return path.stem, content


def _parse_excel(path: Path) -> tuple[str, str]:
    workbook = load_workbook(filename=path, read_only=True, data_only=True)
    sheets: list[str] = []

    for sheet in workbook.worksheets:
        rows: list[str] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))
        if rows:
            sheets.append(f"[sheet {sheet.title}]\n" + "\n".join(rows))

    content = "\n\n".join(sheets).strip()
    if not content:
        raise ValueError("No extractable text found in Excel workbook.")
    return path.stem, content
